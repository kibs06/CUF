#!/usr/bin/env python3
"""Convert Markdown to Word (.docx) format."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

MD_PATH = Path("docs/SoleVision_Project_Documentation.md")
DOCX_PATH = Path("docs/SoleVision_Project_Documentation.docx")


def parse_md_to_docx(md_text: str) -> Document:
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Configure heading styles
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x3B, 0x23, 0x14)  # Carob Dark
        hs.font.name = 'Calibri'
    
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x3B, 0x23, 0x14)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # Skip separator rows
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is still table
            if i + 1 < len(lines) and '|' in lines[i + 1] and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # Render table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < num_cols:
                                cell = table.cell(ri, ci)
                                cell.text = cell_text
                                for p in cell.paragraphs:
                                    p.style = doc.styles['Normal']
                                    for run in p.runs:
                                        run.font.size = Pt(9)
                    doc.add_paragraph()  # spacing after table
                table_rows = []
                i += 1
                continue
        
        stripped = line.strip()
        
        # Empty lines
        if not stripped:
            i += 1
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue
        
        # Horizontal rules
        if stripped in ('---', '***', '___'):
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Unordered list items
        list_match = re.match(r'^[-*]\s+(.*)', stripped)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold and inline code in list items
            _add_formatted_text(p, text)
            i += 1
            continue
        
        # Ordered list items
        ol_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if ol_match:
            text = ol_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1
    
    return doc


def _add_formatted_text(paragraph, text):
    """Add text with bold, inline code, and link formatting."""
    # Pattern: **bold**, `code`, [text](url)
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x5A, 0x2B)
        elif part.startswith('['):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = RGBColor(0x4E, 0xCD, 0xC4)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def main():
    md_text = MD_PATH.read_text(encoding='utf-8')
    doc = parse_md_to_docx(md_text)
    doc.save(str(DOCX_PATH))
    print(f"Converted: {MD_PATH} -> {DOCX_PATH}")
    print(f"File size: {DOCX_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
